import re
import sys
from pathlib import Path
from typing import Optional

from docx import Document

# Imports app / db depuis backend/
sys.path.append(str(Path(__file__).resolve().parent))
from app import app
from db import db, Minimas, Categorie, Epreuve

# ---- mapping nages -> canonique ----
NAGE_MAP = {
    "NAGE LIBRE": "Nage Libre", "NL": "Nage Libre",
    "DOS": "Dos",
    "BRASSE": "Brasse", "BR": "Brasse",
    "PAPILLON": "Papillon", "PAP": "Papillon",
    "4 NAGES": "4 Nages", "4_NAGES": "4 Nages", "4NAGES": "4 Nages",
}

NULL_TOKENS = {"", "-", "—", "–", "N/A", "NA", "MINIMA", "MINIMAS"}  # ignore aussi en-tête "minima"

def category_from_filename(name: str) -> str:
    stem = Path(name).stem
    low  = stem.lower()
    if low == "tc" or low.startswith("tc"): return "TC"
    if "poussin"  in low: return "Poussin"
    if "minime"   in low: return "Minimes"
    if "benjamin" in low: return "Benjamins"
    if "cadet"    in low: return "Cadets"
    if "junior" in low and "senior" in low: return "Juniors/Seniors"
    return stem.title()

def norm_time(s: str) -> str:
    if s is None: return ""
    s = s.strip().replace("\u00A0", "")
    
    s = re.sub(r"\s*:\s*", ":", s)  # "2 :39.84" -> "2:39.84"
    s = (s.replace("：", ":")
         .replace(" ;", ":").replace(";", ":")
         .replace(",,", ".").replace(",", "."))
    return s.strip()

def is_empty(s: str) -> bool:
    return s.strip().upper() in NULL_TOKENS

def normalize_nage_token(tok: str) -> str:
    t = tok.upper().replace("_", " ")
    t = re.sub(r"\s+", " ", t).strip()
    return NAGE_MAP.get(t, t.title())

# ---- ORM helpers avec dédoublonnage DB ----
def ensure_categorie(nom: str) -> Categorie:
    obj = Categorie.query.filter_by(nom=nom).first()
    if not obj:
        obj = Categorie(nom=nom); db.session.add(obj); db.session.flush()
    return obj

def ensure_epreuve(distance: int, nage_token: str, genre: str, is_relay: bool, legs_count: Optional[int]) -> Epreuve:
    """
    Règles :
    - Si relais → distance = distance PAR RELAIS (ex. 50 pour 4x50) + legs_count (4,10…)
    - Si individuel → legs_count doit être NULL
    - Unicité logique alignée avec ta contrainte :
        UNIQUE (nage, distance, genre, is_relay, legs_count)
    """
    nage = normalize_nage_token(nage_token)
    genre = genre.capitalize()

    q = Epreuve.query.filter_by(
        distance=distance, nage=nage, genre=genre, is_relay=is_relay, legs_count=legs_count
    )
    ep = q.first()
    if not ep:
        ep = Epreuve(distance=distance, nage=nage, genre=genre, is_relay=is_relay, legs_count=legs_count)
        db.session.add(ep)
        db.session.flush()
    else:
        # sécurités de cohérence
        if is_relay and ep.legs_count != legs_count:
            ep.legs_count = legs_count
            db.session.flush()
        if not is_relay and ep.legs_count is not None:
            ep.legs_count = None
            db.session.flush()
    return ep

def upsert_minima(epreuve_id: int, categorie_id: int, temp_min: str):
    rec = Minimas.query.filter_by(epreuve_id=epreuve_id, categorie_id=categorie_id).first()
    if rec:
        rec.temp_min = temp_min   # update si déjà là
    else:
        db.session.add(Minimas(epreuve_id=epreuve_id, categorie_id=categorie_id, temp_min=temp_min))

# ---- Parsers d’intitulé d’épreuve ----
# Variante underscores (ex: 4_x_50_m_NAGE_LIBRE_Dames_Classement)
EVENT_U = re.compile(
    r"""^(?:(?P<relay>\d+)_x_)?      # 4_x_ ou 10_x_ (optionnel)
        (?P<dist>\d+)_m_
        (?P<nage>[A-Za-z0-9_]+)_
        (?P<genre>Dames|Messieurs|Mixte)
        (?:_.*)?$                    # suffixe optionnel
    """, re.IGNORECASE | re.VERBOSE
)

# Variante espaces (ex: 4 x 50 m Nage Libre Dames Classement)
EVENT_S = re.compile(
    r"""^(?:(?P<relay>\d+)\s*[xX]\s*)? # 4 x (optionnel)
        (?P<dist>\d+)\s*m\s+
        (?P<nage>(?:NAGE\s*LIBRE|NL|DOS|BRASSE|BR|PAPILLON|PAP|4\s*NAGES))\s+
        (?P<genre>Dames|Messieurs|Mixte)
        (?:\s+.*)?$                   # suffixe optionnel
    """, re.IGNORECASE | re.VERBOSE
)

def match_event(text: str):
    txt = text.strip()
    m = EVENT_U.match(txt.replace(" ", ""))  # variantes compactes
    if m: return m
    return EVENT_S.match(txt)                # variantes avec espaces

# ---- import principal : tables d'abord (col0=épreuve, col1=temps) ----
# ---- Lecture du .docx (tables 2 colonnes, puis fallback paragraphes) ----
def parse_docx(path: Path):
    doc = Document(str(path))
    rows = []

    # 1) tables (format 2 colonnes : libellé épreuve | valeur minima)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = []
            for cell in row.cells:
                t = "\n".join(p.text for p in cell.paragraphs if p.text.strip()).strip()
                cells.append(t)
            if len(cells) >= 2:
                rows.append((cells[0].strip(), cells[1].strip()))

    # 2) fallback : paragraphes successifs
    if not rows:
        lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        i = 0
        while i < len(lines):
            head = lines[i]
            m = match_event(head)
            if m:
                remainder = head[m.end():].strip()
                time_txt = remainder or (lines[i+1] if i+1 < len(lines) else "")
                rows.append((head, time_txt))
                i += 2
            else:
                i += 1

    return rows


# ---- Import principal (par dossier) ----
def load_all(folder: Path):
    total_ins = 0

    # dédoublonnage en mémoire (inclut legs_count pour différencier 4x50 vs 10x50)
    seen_epreuves = set()   # (distance_par_relais, nage_norm, genre, is_relay, legs_count)
    seen_minimas  = set()   # (distance_par_relais, nage_norm, genre, is_relay, legs_count, categorie_nom)

    for fp in sorted(folder.glob("*.docx")):
        cat_nom = category_from_filename(fp.name)
        cat = ensure_categorie(cat_nom)
        rows = parse_docx(fp)

        ins = 0
        for ev_text, time_text in rows:
            if not ev_text:
                continue
            tnorm = norm_time(time_text)
            if is_empty(tnorm):
                continue  # pas de minima à enregistrer

            m = match_event(ev_text)
            if not m:
                continue  # intitulé non reconnu

            gd      = m.groupdict()
            relay   = gd.get("relay")
            dist    = int(gd["dist"])                 # distance PAR RELAIS (ex: 50)
            nage    = gd["nage"]
            genre   = gd["genre"].capitalize()
            is_rel  = relay is not None
            legs    = int(relay) if relay else None   # 4, 10, … ou None si individuel
            nage_n  = normalize_nage_token(nage)

            # upsert épreuve (distance par relais + legs_count)
            ekey = (dist, nage_n, genre, is_rel, legs)
            if ekey not in seen_epreuves:
                seen_epreuves.add(ekey)
            ep = ensure_epreuve(dist, nage, genre, is_rel, legs)

            # upsert minima (par épreuve & catégorie)
            mkey = (dist, nage_n, genre, is_rel, legs, cat_nom)
            if mkey in seen_minimas:
                continue
            seen_minimas.add(mkey)

            upsert_minima(ep.epreuve_id, cat.categorie_id, tnorm)
            ins += 1
            total_ins += 1

        print(f"[OK] {fp.name}: insérés/mis à jour: {ins} (catégorie={cat_nom})")

    db.session.commit()
    return total_ins

if __name__ == "__main__":
    # Dossier d’entrée par défaut : backend/data/minimas
    data_dir = Path(__file__).parent / "data" / "minimas"
    if not data_dir.exists():
        print(f"Répertoire introuvable: {data_dir}")
        raise SystemExit(1)

    with app.app_context():
        total = load_all(data_dir)
        print(f"Import terminé: {total} minima insérés/mis à jour.")
